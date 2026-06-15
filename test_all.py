import os
import django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'resume_project.settings')
django.setup()

import docx
from screening.resume_parser import extract_text_from_file, parse_resume_text
from screening.ml_models import predict_candidate

# Create a sample docx
doc = docx.Document()
doc.add_paragraph("Alice Smith\nBackend Developer")
doc.add_paragraph("Experience: 4 years of experience as a software engineer.")
doc.add_paragraph("Education: B.Tech in Computer Science.")
doc.add_paragraph("Projects: 5 projects.")
doc.add_paragraph("Skills: Python, Django, SQL, REST APIs, Git.")
doc.save("sample_resume.docx")

class DummyFile:
    def __init__(self, path):
        self.path = path
        self.name = os.path.basename(path)
        self.f = open(path, 'rb')
        
    def read(self, *args, **kwargs):
        return self.f.read(*args, **kwargs)
        
    def seek(self, *args, **kwargs):
        return self.f.seek(*args, **kwargs)
        
    def tell(self, *args, **kwargs):
        return self.f.tell(*args, **kwargs)

for fpath in ['test_resume.pdf', 'sample_resume.txt', 'sample_resume.docx']:
    print(f"--- TESTING {fpath} ---")
    try:
        dummy = DummyFile(fpath)
        text = extract_text_from_file(dummy)
        print("Extracted Text Length:", len(text))
        
        parsed = parse_resume_text(text)
        print("Parsed:", parsed)
        
        res = predict_candidate(
            parsed['skills'], 
            parsed['experience'], 
            parsed['education'], 
            parsed['projects_count']
        )
        print("Prediction:", res['prediction'])
    except Exception as e:
        print(f"Error testing {fpath}: {e}")
    print()
